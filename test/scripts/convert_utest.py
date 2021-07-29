import unittest
import argparse
import os
import docx
import logging

from scripts.convert import Convert


class TestGetValidFileTypes(unittest.TestCase):
    def test_get_valid_file_types_idml(self) -> None:
        Convert.args = argparse.Namespace(outputfiletype="idml")
        want_list = ["idml"]

        got_list = Convert.get_valid_file_types(Convert)
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_blank_filetype(self) -> None:
        Convert.args = argparse.Namespace(outputfiletype="", outputfile="cornucopia_en_template.pdf")
        want_list = ["pdf"]

        got_list = Convert.get_valid_file_types(Convert)
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_revert_default(self) -> None:
        Convert.args = argparse.Namespace(outputfiletype="", outputfile="cornucopia_en_template")
        want_list = ["docx"]

        got_list = Convert.get_valid_file_types(Convert)
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_all_with_pdf(self) -> None:
        Convert.args = argparse.Namespace(outputfiletype="all")
        Convert.can_convert_to_pdf = True
        want_list = ["docx", "idml", "pdf"]

        got_list = Convert.get_valid_file_types(Convert)
        got_list.sort()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_all_without_pdf(self) -> None:
        Convert.args = argparse.Namespace(outputfiletype="all")
        Convert.can_convert_to_pdf = False
        want_list = ["docx", "idml"]

        got_list = Convert.get_valid_file_types(Convert)
        got_list.sort()
        self.assertListEqual(want_list, got_list)

    def test_get_valid_file_types_pdf_without_pdf_ability(self) -> None:
        Convert.args = argparse.Namespace(outputfiletype="pdf")
        Convert.can_convert_to_pdf = False
        want_list = []

        logging.getLogger().setLevel(logging.CRITICAL)
        got_list = Convert.get_valid_file_types(Convert)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertListEqual(want_list, got_list)


class TestGetValidLanguagesChoices(unittest.TestCase):
    def test_get_valid_language_choices_fr(self) -> None:
        Convert.args = argparse.Namespace(language="fr")
        want_language = ["fr"]

        got_language = Convert.get_valid_language_choices(Convert)
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_blank(self) -> None:
        Convert.args = argparse.Namespace(language="")
        want_language = ["en"]

        got_language = Convert.get_valid_language_choices(Convert)
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_all(self) -> None:
        Convert.args = argparse.Namespace(language="all")
        want_language = Convert.LANGUAGE_CHOICES
        want_language.remove("template")
        want_language.remove("all")

        got_language = Convert.get_valid_language_choices(Convert)
        self.assertListEqual(want_language, got_language)

    def test_get_valid_language_choices_template(self) -> None:
        Convert.args = argparse.Namespace(language="template")
        want_language = ["template"]

        got_language = Convert.get_valid_language_choices(Convert)
        self.assertListEqual(want_language, got_language)


class TestSetCanConvertToPdf(unittest.TestCase):
    def test_set_can_convert_to_pdf(self) -> None:
        Convert.can_convert_to_pdf = None
        want_can_convert_in = [True, False]

        Convert.set_can_convert_to_pdf(Convert)
        got_can_convert = Convert.can_convert_to_pdf
        self.assertIn(got_can_convert, want_can_convert_in)


class TestSortKeysLongestToShortest(unittest.TestCase):
    def test_sort_keys_longest_to_shortest_success(self) -> None:
        source_data = {
            "key shortest": "value1",
            "key ------------- longest": "value2",
            "key ------ middle": "value3",
        }
        want_data = [
            ("key ------------- longest", "value2"),
            ("key ------ middle", "value3"),
            ("key shortest", "value1"),
        ]

        got_data = Convert.sort_keys_longest_to_shortest(source_data)
        self.assertEqual(want_data, got_data)


class TestSetLogging(unittest.TestCase):
    def test_set_logging_default(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        want_logging_level = logging.INFO

        Convert.set_logging(Convert)
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)

    def test_set_logging_true(self) -> None:
        Convert.args = argparse.Namespace(debug=True)
        want_logging_level = logging.DEBUG

        Convert.set_logging(Convert)
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)

    def test_set_logging_info(self) -> None:
        Convert.args = argparse.Namespace(debug=None)
        want_logging_level = logging.INFO

        Convert.set_logging(Convert)
        got_logging_level = logging.getLogger().level
        self.assertEqual(want_logging_level, got_logging_level)


class TestParseArguments(unittest.TestCase):
    def test_successfully_parsing_arguments_short_form_basic(self) -> None:
        input_args = ["-t", "idml"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="en",
            debug=False,
        )

        got_args = Convert.parse_arguments(Convert, input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)

    def test_successfully_parsing_arguments_short_form_language(self) -> None:
        input_args = ["-t", "idml", "-l", "fr"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="fr",
            debug=False,
        )

        got_args = Convert.parse_arguments(Convert, input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)

    def test_successfully_parsing_arguments_long_form_basic(self) -> None:
        input_args = ["--outputfiletype", "idml"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="en",
            debug=False,
        )

        got_args = Convert.parse_arguments(Convert, input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)

    def test_successfully_parsing_arguments_long_form_language(self) -> None:
        input_args = ["--outputfiletype", "idml", "--language", "fr"]
        want_args = argparse.Namespace(
            inputfile="",
            outputfiletype="idml",
            outputfile="",
            language="fr",
            debug=False,
        )

        got_args = Convert.parse_arguments(Convert, input_args)
        self.maxDiff = None
        self.assertEqual(got_args, want_args)


class TestMakeTemplate(unittest.TestCase):
    def test_set_making_template_true(self) -> None:
        Convert.args = argparse.Namespace(language="template")

        Convert.set_making_template(Convert)
        result = Convert.making_template
        self.assertTrue(result)

    def test_set_making_template_false(self) -> None:
        Convert.args = argparse.Namespace(language="en")

        Convert.set_making_template(Convert)
        result = Convert.making_template
        self.assertFalse(result)

    def test_set_making_template_empty(self) -> None:
        Convert.args = argparse.Namespace()

        Convert.set_making_template(Convert)
        result = Convert.making_template
        self.assertFalse(result)


class TestGetFilesFromOfType(unittest.TestCase):
    def test_get_files_from_of_type_source_yaml_files(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        path = Convert.BASE_PATH + "/test/test_files/"
        ext = "yaml"
        want_files = ["ecommerce-cards-1.21-en.yaml", "ecommerce-mappings-1.2.yaml"]
        want_count = len(want_files)

        got_files = Convert.get_files_from_of_type(Convert, path, ext)
        self.assertEqual(len(got_files), want_count)
        got_files = list(os.path.basename(f) for f in got_files)
        self.assertListEqual(got_files, want_files)

    def test_get_files_from_of_type_source_docx_files(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        path = Convert.BASE_PATH + "/test/test_files/"
        ext = "docx"
        want_files = ["owasp_cornucopia_edition_lang_ver_template.docx"]

        got_files = Convert.get_files_from_of_type(Convert, path, ext)
        got_files = list(os.path.basename(f) for f in got_files)
        self.assertListEqual(got_files, want_files)
        for f in want_files:
            self.assertIn(f, got_files)

    def test_get_files_from_of_type_source_empty_list(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        path = Convert.BASE_PATH + "/test/test_files/"
        ext = "ext"
        want_files = []

        logging.getLogger().setLevel(logging.CRITICAL)
        got_files = Convert.get_files_from_of_type(Convert, path, ext)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertListEqual(got_files, want_files)


class TestGetDocxDocument(unittest.TestCase):
    def test_get_docx_document_success(self) -> None:
        file = Convert.BASE_PATH + "/test/test_files/owasp_cornucopia_edition_lang_ver_template.docx"
        want_type = docx.document.Document
        want_len_paragraphs = 36

        got_file = Convert.get_docx_document(file)
        got_type = type(got_file)
        self.assertEqual(want_type, got_type)
        got_len_paragraphs = len(got_file.paragraphs)
        self.assertEqual(want_len_paragraphs, got_len_paragraphs)

    def test_get_docx_document_failure(self) -> None:
        file = Convert.BASE_PATH + "/test/test_files/owasp_cornucopia_edition_lang_ver_template.d"
        want_type = docx.document.Document
        want_len_paragraphs = 0

        logging.getLogger().setLevel(logging.CRITICAL)
        got_file = Convert.get_docx_document(file)
        logging.getLogger().setLevel(logging.ERROR)
        self.assertIsInstance(got_file, want_type)
        got_len_paragraphs = len(got_file.paragraphs)
        self.assertEqual(want_len_paragraphs, got_len_paragraphs)


class TestGetTagForSuitName(unittest.TestCase):
    def test_get_tag_for_suit_name_ve(self) -> None:
        Convert.making_template = False
        suit = {"name": "Data validation & encoding", "cards": []}
        suit_tag = "VE"
        want_tag_data = {"${VE_suit}": suit["name"]}

        got_tag_data = Convert.get_tag_for_suit_name(suit, suit_tag)
        self.assertDictEqual(want_tag_data, got_tag_data)

    def test_get_tag_for_suit_name_ve_template(self) -> None:
        Convert.making_template = True
        suit = {"name": "Data validation & encoding", "cards": []}
        suit_tag = "VE"
        want_tag_data = {suit["name"]: "${VE_suit}"}

        got_tag_data = Convert.get_tag_for_suit_name(suit, suit_tag)
        self.assertDictEqual(want_tag_data, got_tag_data)


class TestGetReplacementDict(unittest.TestCase):
    test_data = {
        "meta": {"edition": "ecommerce", "component": "cards", "language": "EN", "version": "1.21"},
        "suits": [
            {
                "name": "Data validation & encoding",
                "cards": [
                    {
                        "value": "A",
                        "desc": "You have invented a new attack against Data Validation",
                        "misc": "Read more about this topic in OWASP's free Cheat Sheets",
                    },
                    {"value": "2", "desc": "Brian can gather information about the underlying configurations"},
                ],
            },
            {
                "name": "Authentication",
                "cards": [
                    {
                        "value": "A",
                        "desc": "You have invented a new attack against Authentication",
                        "misc": "Read more about this topic in OWASP's free Cheat Sheet",
                    },
                    {"value": "2", "desc": "James can undertake authentication functions without"},
                ],
            },
        ],
    }

    def test_get_replacement_dict_success(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        logging.basicConfig(level=logging.ERROR)
        input_data = self.test_data
        want_type = dict
        want_length = 8
        want_data = {
            "${VE_suit}": "Data validation & encoding",
            "${VE_VEA_desc}": "You have invented a new attack against Data Validation",
            "${VE_VEA_misc}": "Read more about this topic in OWASP's free Cheat Sheets",
            "${VE_VE2_desc}": "Brian can gather information about the underlying configurations",
            "${AT_suit}": "Authentication",
            "${AT_ATA_desc}": "You have invented a new attack against Authentication",
            "${AT_ATA_misc}": "Read more about this topic in OWASP's free Cheat Sheet",
            "${AT_AT2_desc}": "James can undertake authentication functions without",
        }

        got_data = Convert.get_replacement_dict(Convert, input_data)
        self.assertIsInstance(got_data, want_type)
        self.assertEqual(len(got_data), want_length)
        self.assertDictEqual(got_data, want_data)

    def test_get_replacement_dict_template(self) -> None:
        Convert.making_template = True
        logging.basicConfig(level=logging.ERROR)
        input_data = self.test_data
        want_type = dict
        want_length = 8
        want_data = {
            "Data validation & encoding": "${VE_suit}",
            "You have invented a new attack against Data Validation": "${VE_VEA_desc}",
            "Read more about this topic in OWASP's free Cheat Sheets": "${VE_VEA_misc}",
            "Brian can gather information about the underlying configurations": "${VE_VE2_desc}",
            "Authentication": "${AT_suit}",
            "You have invented a new attack against Authentication": "${AT_ATA_desc}",
            "Read more about this topic in OWASP's free Cheat Sheet": "${AT_ATA_misc}",
            "James can undertake authentication functions without": "${AT_AT2_desc}",
        }

        got_data = Convert.get_replacement_dict(Convert, input_data)
        self.assertIsInstance(got_data, want_type)
        self.assertEqual(len(got_data), want_length)
        self.assertDictEqual(got_data, want_data)


class TestGetCheckFixFileExtension(unittest.TestCase):
    def test_get_check_fix_file_extension_no_extension(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        input_filename = "hello"
        input_extension = ".docx"
        want_filename = "hello.docx"

        got_filename = Convert.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_wrong_extension(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        input_filename = "hello.docx"
        input_extension = ".pdf"
        want_filename = "hello.pdf"

        got_filename = Convert.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)

    def test_get_check_fix_file_extension_correct_extension(self) -> None:
        Convert.args = argparse.Namespace(debug=False)
        input_filename = "hello.docx"
        input_extension = ".docx"
        want_filename = "hello.docx"

        got_filename = Convert.check_fix_file_extension(input_filename, input_extension)
        self.assertEqual(want_filename, got_filename)
